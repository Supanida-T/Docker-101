from docx import Document

# เปิดไฟล์ docx
doc = Document('template.docx')

# พจนานุกรมเก็บคำที่ต้องการเปลี่ยน: {"คำเก่า": "คำใหม่"}
replacements = {
    "ศุภมงคล": "นริณ",
    "ใจเทียมศักดิ์": "กำแพงแสน",
    "จารแบน": "กำแพงแสน",
    "ธีรวัฒน์ฯ": "สมชายฯ",
    "ปิ่นวรางค์": "วรารัตน์",
    "กลิ่นหวล": "งามเลิศ",
    "วิชัย": "ชัย",
    "ภัสโรวัฒนา": "คำแสง",
    "ณัฏฐพัชร์": "นราวิช",
    "งามประดิษฐ์": "ธีรนันท์",
    "๐๙๕ ๖๑๐ ๑๗๘๖": "๐๘๓ ๕๒๖ ๑๙๐",
    "๓๕/๒": "๒/๑",
    "๓": "๘",
    "๑-๒๖๐๔-๐๑๑๕๐-๓๙-๕": "๑-๕๐๒๔-๔๕๖๗-๘๙-๐"


    # เพิ่มข้อมูลที่ต้องการเปลี่ยนได้ที่นี่
}

# ลูปผ่านแต่ละพารากราฟและแก้ไขเนื้อหาที่มีคำเก่า
for para in doc.paragraphs:
    for run in para.runs:
        for old_name, new_name in replacements.items():
            if old_name in run.text:
                run.text = run.text.replace(old_name, new_name)

# ลูปผ่านแต่ละเซลล์ในตารางเพื่อแก้ไขเนื้อหาที่มีคำเก่า
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for old_name, new_name in replacements.items():
                        if old_name in run.text:
                            run.text = run.text.replace(old_name, new_name)

# บันทึกไฟล์ใหม่
doc.save('outputบันทึกคำให้การผู้ต้องหา.docx')

print("เปลี่ยนหลายข้อมูลในเอกสารเสร็จเรียบร้อยแล้ว!")
