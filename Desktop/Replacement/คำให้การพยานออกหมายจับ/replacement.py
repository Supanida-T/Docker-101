from docx import Document

# เปิดไฟล์ docx
doc = Document('คำให้การพยานออกหมายจับ.docx')

# พจนานุกรมเก็บคำที่ต้องการเปลี่ยน: {"คำเก่า": "คำใหม่"}
replacements = {
    "ธีรวัฒน์": "นริณ",
    "จานแบน": "กำแพงแสน",
    "จารแบน": "กำแพงแสน",
    "ธีรวัฒน์ฯ": "นริณฯ",
    "ปิ่นวรางค์": "วรารัตน์",
    "กลิ่นหวล": "งามเลิศ",
    "วิภาวรรณ": "สุดา",
    "พุฒนอก": "คำแสง",
    "วรวุฒิ": "สุชาติ",
    "วรวุฒิฯ ": "สุชาติฯ",
    "ตูน": "ต้า",
    

    # เพิ่มข้อมูลที่ต้องการเปลี่ยนได้ที่นี่
}

# ลูปผ่านแต่ละพารากราฟและแก้ไขเนื้อหาที่มีคำเก่า
for para in doc.paragraphs:
    for run in para.runs:
        for old_name, new_name in replacements.items():
            if old_name in run.text:
                run.text = run.text.replace(old_name, new_name)

# ลูปผ่านแต่ละเซลล์ในตารางเพื่อแก้ไขเนื้อหาที่มีคำเก่า
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for old_name, new_name in replacements.items():
                        if old_name in run.text:
                            run.text = run.text.replace(old_name, new_name)

# บันทึกไฟล์ใหม่
doc.save('outputคำให้การพยานออกหมายจับ.docx')

print("เปลี่ยนหลายข้อมูลในเอกสารเสร็จเรียบร้อยแล้ว!")
