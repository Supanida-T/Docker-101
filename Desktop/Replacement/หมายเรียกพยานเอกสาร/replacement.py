from docx import Document

# เปิดไฟล์ docx
doc = Document('หมายเรียกพยานเอกสารข้อมูลโทรศัพท์.docx')

# พจนานุกรมเก็บคำที่ต้องการเปลี่ยน: {"คำเก่า": "คำใหม่"}
replacements = {
    "4 เดือน ธันวาคม": "2 เดือน มกราคม",
    "2566": "2562",
    "0870672337": "0812345678",
    "0817538561": "0823456789",
    "วันที่ 4 เดือน กันยายน": "วันที่ 2 เดือน กันยายน",
    "สุรศักดิ์": "อธิศัก",
    "หญีตบึ้ง": "บุญมา",
    "ก่อเกียรติ": "ธิป",
    "วุฒิจำนงค์": "ทวีนา",
    "sub1.atpd@cib.go.th": "s1@cib.co.th",
        
    
    

    # เพิ่มข้อมูลที่ต้องการเปลี่ยนได้ที่นี่
}

# ลูปผ่านแต่ละพารากราฟและแก้ไขเนื้อหาที่มีคำเก่า
for para in doc.paragraphs:
    for run in para.runs:
        for old_name, new_name in replacements.items():
            if old_name in run.text:
                run.text = run.text.replace(old_name, new_name)

# ลูปผ่านแต่ละเซลล์ในตารางเพื่อแก้ไขเนื้อหาที่มีคำเก่า
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for old_name, new_name in replacements.items():
                        if old_name in run.text:
                            run.text = run.text.replace(old_name, new_name)

# บันทึกไฟล์ใหม่
doc.save('outputหมายเรียกพยานเอกสารข้อมูลโทรศัพท์.docx')

print("เปลี่ยนหลายข้อมูลในเอกสารเสร็จเรียบร้อยแล้ว!")
