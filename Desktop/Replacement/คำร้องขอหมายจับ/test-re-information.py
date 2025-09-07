from docx import Document

# เปิดไฟล์ docx
doc = Document('คำร้องขอหมายจับ.docx')

# ชื่อเก่าที่ต้องการเปลี่ยนและชื่อใหม่
old_name = "ธีรวัฒน์ จานแบน"
new_name = "สมหมาย กำแพงแสน"

# ลูปผ่านแต่ละพารากราฟเพื่อหาชื่อและแทนที่
for para in doc.paragraphs:
    if old_name in para.text:
        para.text = para.text.replace(old_name, new_name)

# ลูปผ่านแต่ละเซลล์ในตาราง (ถ้ามีตารางในเอกสาร)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if old_name in cell.text:
                cell.text = cell.text.replace(old_name, new_name)

# บันทึกไฟล์ใหม่
doc.save('คำร้องขอหมายจับ.docx')

print("เปลี่ยนชื่อในเอกสารเสร็จเรียบร้อยแล้ว!")