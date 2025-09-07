from docx import Document

# เปิดไฟล์ docx
doc = Document('คำร้องขอหมายจับ.docx')

# ชื่อเก่าที่ต้องการเปลี่ยนและชื่อใหม่
old_name = "ธีรวัฒน์ จานแบน"
new_name = "สมชาย กำแพงแสน"

# ลูปผ่านแต่ละพารากราฟและแก้ไขเนื้อหาที่มีชื่อเก่า
for para in doc.paragraphs:
    for run in para.runs:  # ลูปผ่าน runs เพื่อรักษาฟอร์แมต
        if old_name in run.text:
            run.text = run.text.replace(old_name, new_name)

# ลูปผ่านแต่ละเซลล์ในตารางเพื่อแก้ไขเนื้อหาที่มีชื่อเก่า
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:  # รักษาฟอร์แมตในตาราง
                    if old_name in run.text:
                        run.text = run.text.replace(old_name, new_name)

# บันทึกไฟล์ใหม่
doc.save('output.docx')

print("เปลี่ยนชื่อในเอกสารเสร็จเรียบร้อยแล้ว พร้อมรักษาฟอร์แมตและเส้นตาราง!")
